from __future__ import annotations

from pathlib import Path
import subprocess
import zipfile

from docx import Document
import pytest

import app.services.competition_legacy_doc_converter as converter_module
from app.services.competition_legacy_doc_converter import (
    CompetitionLegacyDocConversionError,
    convert_legacy_doc_to_docx,
    find_libreoffice_executable,
)


def _fake_soffice(
    tmp_path: Path,
) -> Path:
    executable = (
        tmp_path
        / "soffice.exe"
    )

    executable.write_bytes(
        b"fake"
    )

    return executable


def test_find_libreoffice_uses_explicit_path(
    tmp_path: Path,
) -> None:
    executable = _fake_soffice(
        tmp_path
    )

    resolved = (
        find_libreoffice_executable(
            explicit_path=executable,
        )
    )

    assert resolved == executable.resolve()


def test_convert_legacy_doc_builds_headless_command(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "legacy.doc"

    original_bytes = (
        bytes.fromhex(
            "D0CF11E0A1B11AE1"
        )
        + b"legacy"
    )

    source.write_bytes(
        original_bytes
    )

    executable = _fake_soffice(
        tmp_path
    )

    output_directory = (
        tmp_path
        / "converted"
    )

    captured: dict[
        str,
        object,
    ] = {}

    def fake_run(
        command,
        **kwargs,
    ):
        captured["command"] = command
        captured["kwargs"] = kwargs

        document = Document()
        document.add_paragraph(
            "转换后的监管制度正文。"
        )

        table = document.add_table(
            rows=2,
            cols=2,
        )

        table.cell(0, 0).text = "指标"
        table.cell(0, 1).text = "要求"
        table.cell(1, 0).text = "资本"
        table.cell(1, 1).text = "符合要求"

        destination = (
            output_directory
            / "legacy.docx"
        )

        document.save(
            destination
        )

        return subprocess.CompletedProcess(
            args=command,
            returncode=0,
            stdout="converted",
            stderr="",
        )

    monkeypatch.setattr(
        converter_module.subprocess,
        "run",
        fake_run,
    )

    result = convert_legacy_doc_to_docx(
        source_path=source,
        output_directory=output_directory,
        soffice_path=executable,
        timeout_seconds=12,
    )

    assert result == (
        output_directory
        / "legacy.docx"
    )

    assert result.is_file()
    assert zipfile.is_zipfile(result)

    # 原始 DOC 不允许被修改。
    assert (
        source.read_bytes()
        == original_bytes
    )

    command = captured["command"]
    kwargs = captured["kwargs"]

    assert "--headless" in command
    assert "--convert-to" in command
    assert "docx" in command
    assert "--outdir" in command

    assert kwargs["timeout"] == 12
    assert kwargs["check"] is False
    assert kwargs["capture_output"] is True


def test_convert_legacy_doc_requires_output_file(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "legacy.doc"

    source.write_bytes(
        bytes.fromhex(
            "D0CF11E0A1B11AE1"
        )
        + b"legacy"
    )

    executable = _fake_soffice(
        tmp_path
    )

    def fake_run(
        command,
        **kwargs,
    ):
        return subprocess.CompletedProcess(
            args=command,
            returncode=0,
            stdout="",
            stderr="",
        )

    monkeypatch.setattr(
        converter_module.subprocess,
        "run",
        fake_run,
    )

    with pytest.raises(
        CompetitionLegacyDocConversionError,
        match="未生成预期 DOCX",
    ):
        convert_legacy_doc_to_docx(
            source_path=source,
            output_directory=(
                tmp_path
                / "converted"
            ),
            soffice_path=executable,
        )