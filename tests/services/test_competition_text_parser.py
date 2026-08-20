from __future__ import annotations

from pathlib import Path

import pymupdf
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.schemas.competition import (
    CompetitionQuestion,
    CompetitionSourceRecord,
)
from app.services.competition_text_parser import (
    CompetitionUnsupportedTextFormatError,
    parse_competition_text_document,
)


def _source_record(
    *,
    path: Path,
    source_type: str,
) -> CompetitionSourceRecord:
    return CompetitionSourceRecord(
        source_id=(
            "src_0123456789abcdef"
        ),
        source_type=source_type,
        actual_filename=path.name,
        relative_path=path.name,
        extension=(
            path.suffix.casefold()
        ),
        size_bytes=(
            path.stat().st_size
        ),
    )


def _question(
    *,
    source_type: str,
    file_label: str,
) -> CompetitionQuestion:
    return CompetitionQuestion(
        case_id="Q001",
        source_type=source_type,
        qa_type="单事实检索",
        question=(
            "根据附件，"
            "以下哪项说法正确？"
        ),
        option_a="A",
        option_b="B",
        option_c="C",
        option_d="D",
        source_title="测试监管文件",
        file_label=file_label,
    )


def test_parse_pdf_creates_page_blocks(
    tmp_path: Path,
) -> None:
    attachments = (
        tmp_path
        / "attachments"
    )

    attachments.mkdir()

    pdf_path = (
        attachments
        / "test.pdf"
    )

    document = pymupdf.open()

    page1 = document.new_page()

    page1.insert_text(
        (72, 72),
        "Article 1 Bank risk management.",
    )

    page2 = document.new_page()

    page2.insert_text(
        (72, 72),
        "Article 2 Internal control.",
    )

    document.save(
        pdf_path
    )

    document.close()

    source = _source_record(
        path=pdf_path,
        source_type="pdf",
    )

    question = _question(
        source_type="pdf",
        file_label="test.pdf",
    )

    parsed = (
        parse_competition_text_document(
            question=question,
            source=source,
            attachments_root=attachments,
        )
    )

    assert (
        parsed.source.source_type
        == "pdf"
    )

    assert len(
        parsed.blocks
    ) == 2

    assert (
        parsed.blocks[0]
        .block_type
        == "page_text"
    )

    assert (
        parsed.blocks[0].page
        == 1
    )

    assert (
        parsed.blocks[1].page
        == 2
    )

    assert (
        "Article 1"
        in parsed.blocks[0].text
    )

    assert (
        parsed.source.sha256
        is not None
    )


def test_parse_docx_preserves_paragraph_table_order(
    tmp_path: Path,
) -> None:
    attachments = (
        tmp_path
        / "attachments"
    )

    attachments.mkdir()

    docx_path = (
        attachments
        / "test.docx"
    )

    document = Document()

    # ========================================================
    # Heading 1
    #
    # python-docx 默认：
    # style_name = Heading 1
    # outline_level = 0
    # ========================================================

    document.add_heading(
        "第一章 总则",
        level=1,
    )

    # 故意加入空段落。
    #
    # Parser 会跳过空文本 Block，
    # 但 paragraph_index 仍然应该递增。
    document.add_paragraph(
        ""
    )

    # ========================================================
    # Table
    # ========================================================

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(
        0,
        0,
    ).text = "指标"

    table.cell(
        0,
        1,
    ).text = "要求"

    table.cell(
        1,
        0,
    ).text = "资本充足率"

    table.cell(
        1,
        1,
    ).text = "不得低于规定标准"

    # ========================================================
    # 普通 Paragraph
    # ========================================================

    document.add_paragraph(
        "第二条 商业银行应当"
        "建立风险管理制度。"
    )

    document.save(
        docx_path
    )

    source = _source_record(
        path=docx_path,
        source_type="word",
    )

    question = _question(
        source_type="word",
        file_label="test.docx",
    )

    parsed = (
        parse_competition_text_document(
            question=question,
            source=source,
            attachments_root=attachments,
        )
    )

    # ========================================================
    # Block 顺序：
    #
    # 0 Heading Paragraph
    # 1 Table
    # 2 Normal Paragraph
    #
    # 空 Paragraph 不产生 Block。
    # ========================================================

    assert [
        block.block_type
        for block
        in parsed.blocks
    ] == [
        "paragraph",
        "table",
        "paragraph",
    ]

    first_paragraph = (
        parsed.blocks[0]
    )

    table_block = (
        parsed.blocks[1]
    )

    second_paragraph = (
        parsed.blocks[2]
    )

    # ========================================================
    # First Paragraph
    # ========================================================

    assert (
        first_paragraph
        .paragraph_index
        == 0
    )

    assert (
        first_paragraph.style_name
        == "Heading 1"
    )

    assert (
        first_paragraph.outline_level
        == 0
    )

    # ========================================================
    # Second Paragraph
    #
    # paragraph_index=1 是空段落，
    # 虽然它没有生成 Block，
    # 但仍然占据原始 Word paragraph index。
    # ========================================================

    assert (
        second_paragraph
        .paragraph_index
        == 2
    )

    assert (
        second_paragraph.style_name
        == "Normal"
    )

    assert (
        second_paragraph.outline_level
        is None
    )

    # ========================================================
    # Table
    # ========================================================

    assert (
        table_block.table_index
        == 0
    )

    assert (
        table_block.table_rows
        == (
            (
                "指标",
                "要求",
            ),
            (
                "资本充足率",
                "不得低于规定标准",
            ),
        )
    )

    assert (
        "资本充足率"
        in table_block.text
    )

    assert (
        "不得低于规定标准"
        in table_block.text
    )

    # Table 不应该携带 Paragraph Style。
    assert (
        table_block.style_name
        is None
    )

    assert (
        table_block.outline_level
        is None
    )


def test_parser_rejects_legacy_doc(
    tmp_path: Path,
) -> None:
    attachments = (
        tmp_path
        / "attachments"
    )

    attachments.mkdir()

    doc_path = (
        attachments
        / "legacy.doc"
    )

    # OLE Compound File signature。
    doc_path.write_bytes(
        bytes.fromhex(
            "D0CF11E0A1B11AE1"
        )
        + b"test"
    )

    source = _source_record(
        path=doc_path,
        source_type="word",
    )

    question = _question(
        source_type="word",
        file_label="legacy.doc",
    )

    with pytest.raises(
        CompetitionUnsupportedTextFormatError
    ):
        parse_competition_text_document(
            question=question,
            source=source,
            attachments_root=attachments,
        )

def test_parse_docx_preserves_direct_outline_level(
    tmp_path: Path,
) -> None:
    attachments = (
        tmp_path
        / "attachments"
    )

    attachments.mkdir()

    docx_path = (
        attachments
        / "outline.docx"
    )

    document = Document()

    paragraph = (
        document.add_paragraph(
            "信用风险"
        )
    )

    # 保持 Normal Style，
    # 但直接设置 Word outline level 1。
    ppr = (
        paragraph._p
        .get_or_add_pPr()
    )

    outline = OxmlElement(
        "w:outlineLvl"
    )

    outline.set(
        qn("w:val"),
        "1",
    )

    ppr.append(
        outline
    )

    document.save(
        docx_path
    )

    source = _source_record(
        path=docx_path,
        source_type="word",
    )

    question = _question(
        source_type="word",
        file_label="outline.docx",
    )

    parsed = (
        parse_competition_text_document(
            question=question,
            source=source,
            attachments_root=attachments,
        )
    )

    block = parsed.blocks[0]

    assert (
        block.style_name
        == "Normal"
    )

    assert (
        block.outline_level
        == 1
    )

    assert (
        block.text
        == "信用风险"
    )