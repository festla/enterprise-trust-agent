from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from app.schemas.competition import (
    CompetitionQuestion,
    CompetitionSourceRecord,
)
from app.schemas.competition_text import (
    CompetitionTextBlock,
    CompetitionTextDocument,
)
from app.services.competition_source_catalog import (
    build_competition_knowledge_source,
    resolve_competition_source_path,
)
from app.services.page_parser import (
    normalize_page_text,
)
from tempfile import TemporaryDirectory

from app.services.competition_legacy_doc_converter import (
    convert_legacy_doc_to_docx,
)

class CompetitionTextParserError(
    RuntimeError
):
    """Competition 文本文档解析基础异常。"""


class CompetitionUnsupportedTextFormatError(
    CompetitionTextParserError
):
    """当前 Parser 尚不支持的文本文件格式。"""


class CompetitionPdfOpenError(
    CompetitionTextParserError
):
    """Competition PDF 无法打开。"""


class CompetitionDocxOpenError(
    CompetitionTextParserError
):
    """Competition DOCX 无法打开。"""


class CompetitionEmptyTextDocumentError(
    CompetitionTextParserError
):
    """文档没有产生任何可检索文本 Block。"""


def _build_block_id(
    *,
    doc_id: str,
    block_index: int,
) -> str:
    """
    Block ID 与实际文件版本绑定。

    文件内容变化
        -> SHA256 变化
        -> doc_id 变化
        -> block_id 变化
    """

    return (
        f"block:{doc_id}:"
        f"{block_index:05d}"
    )


def _normalize_word_text(
    value: str,
) -> str:
    """
    Word 与 PDF 共用低风险文本规范化。

    当前直接复用 Week1 的 normalize_page_text，
    不重新实现第二套 normalization。
    """

    return normalize_page_text(
        value
    )


# ============================================================
# PDF
# ============================================================


def _parse_pdf(
    *,
    path: Path,
    question: CompetitionQuestion,
    source: CompetitionSourceRecord,
    attachments_root: Path,
) -> CompetitionTextDocument:
    knowledge_source = (
        build_competition_knowledge_source(
            question=question,
            source=source,
            attachments_root=attachments_root,
        )
    )

    document = None

    try:
        document = pymupdf.open(
            path
        )

    except Exception as exc:
        raise CompetitionPdfOpenError(
            f"PDF 无法打开: {path}"
        ) from exc

    try:
        if not document.is_pdf:
            raise CompetitionPdfOpenError(
                f"文件不是有效 PDF: {path}"
            )

        if document.needs_pass:
            raise CompetitionPdfOpenError(
                f"PDF 需要密码: {path}"
            )

        blocks: list[
            CompetitionTextBlock
        ] = []

        for page_index in range(
            document.page_count
        ):
            page = document.load_page(
                page_index
            )

            # =================================================
            # 与旧 page_parser 保持一致：
            # 先进行 PyMuPDF plain-text extraction，
            # 再调用已有 normalize_page_text。
            #
            # 当前不主动切换 sort=True，
            # 避免在还没有 Dev failure 的情况下
            # 改变旧解析行为。
            # =================================================

            raw_text = page.get_text(
                "text"
            )

            if not isinstance(
                raw_text,
                str,
            ):
                raise CompetitionTextParserError(
                    "PyMuPDF 页面文本必须为字符串"
                )

            normalized_text = (
                normalize_page_text(
                    raw_text
                )
            )

            # 空页不产生 Retrieval Block。
            #
            # 当前 QA-used PDF audit 中没有空文本页，
            # 但 Parser 本身仍允许普通 PDF 出现空页。
            if not normalized_text:
                continue

            block_index = len(
                blocks
            )

            blocks.append(
                CompetitionTextBlock(
                    block_id=_build_block_id(
                        doc_id=(
                            knowledge_source
                            .doc_id
                        ),
                        block_index=(
                            block_index
                        ),
                    ),
                    source_id=(
                        knowledge_source
                        .source_id
                    ),
                    doc_id=(
                        knowledge_source
                        .doc_id
                    ),
                    source_type="pdf",
                    block_index=(
                        block_index
                    ),
                    block_type="page_text",
                    text=normalized_text,

                    # PDF 页码对用户采用 1-based。
                    page=(
                        page_index + 1
                    ),
                )
            )

        if not blocks:
            raise (
                CompetitionEmptyTextDocumentError(
                    "PDF 没有产生任何"
                    "可检索文本 Block"
                )
            )

        return CompetitionTextDocument(
            source=knowledge_source,
            blocks=tuple(
                blocks
            ),
        )

    finally:
        if document is not None:
            document.close()


# ============================================================
# DOCX
# ============================================================
def _paragraph_style_name(
    paragraph: Paragraph,
) -> str | None:
    """
    保留 Word 原始 Paragraph Style。

    Style 本身不是最终标题判断依据，
    只是一个结构信号。
    """

    style = paragraph.style

    if style is None:
        return None

    name = (
        style.name
        or style.style_id
    )

    if not name:
        return None

    return name.strip()


def _read_outline_level_from_ppr(
    ppr,
) -> int | None:
    """
    从 OOXML paragraph properties 中读取：

        <w:outlineLvl w:val="1"/>

    """
    if ppr is None:
        return None

    outline = ppr.find(
        qn("w:outlineLvl")
    )

    if outline is None:
        return None

    value = outline.get(
        qn("w:val")
    )

    if value is None:
        return None

    try:
        return int(value)
    except ValueError:
        return None


def _effective_outline_level(
    paragraph: Paragraph,
) -> int | None:
    """
    获取 Paragraph 的有效 Word 大纲层级。

    优先级：

        direct paragraph formatting
                ↓
        paragraph style
                ↓
        base style chain

    这可以正确处理真实比赛文件中的：

        Normal@level0
        Normal@level1
        Normal@level2
    """

    # ========================================================
    # 1. Paragraph 自己设置的 outline level
    # ========================================================

    direct_level = (
        _read_outline_level_from_ppr(
            paragraph._p.pPr
        )
    )

    if direct_level is not None:
        return direct_level

    # ========================================================
    # 2. Style / Base Style
    # ========================================================

    style = paragraph.style

    visited: set[str] = set()

    while style is not None:
        style_id = (
            style.style_id
            or str(id(style))
        )

        if style_id in visited:
            break

        visited.add(
            style_id
        )

        style_level = (
            _read_outline_level_from_ppr(
                style._element.pPr
            )
        )

        if style_level is not None:
            return style_level

        style = style.base_style

    return None

def _extract_table_rows(
    table: Table,
) -> tuple[
    tuple[str, ...],
    ...,
]:
    """
    保留DOCX表格的二维网格结构。

    Word合并单元格会被python-docx重复映射到
    row.cells中的多个网格位置。

    处理规则：

    1. 合并单元格首次出现时保留文本；
    2. 后续重复网格位置写入空字符串；
    3. 不改变原始行数和列数；
    4. 独立单元格即使文本相同也不会被去重。
    """

    rows: list[
        tuple[str, ...]
    ] = []

    # 保存已经出现过的底层XML Cell。
    #
    # 同一个合并单元格在row.cells中会返回
    # 相同的<w:tc>对象。
    seen_table_cells: set[
        object
    ] = set()

    for row in table.rows:
        row_values: list[str] = []

        for cell in row.cells:
            xml_cell = cell._tc

            if xml_cell in seen_table_cells:
                # 保留网格位置，
                # 但不重复写入合并单元格文本。
                row_values.append("")
                continue

            seen_table_cells.add(
                xml_cell
            )

            row_values.append(
                _normalize_word_text(
                    cell.text
                )
            )

        if not row_values:
            continue

        rows.append(
            tuple(row_values)
        )

    return tuple(rows)


def _table_rows_to_text(
    rows: tuple[
        tuple[str, ...],
        ...,
    ],
) -> str:
    """
    将二维表格同时生成一个适合 BM25 / Dense
    使用的文本表示。

    原始二维结构仍然保存在 table_rows，
    并没有被丢弃。
    """

    lines = []

    for row in rows:
        line = "\t".join(
            row
        ).strip()

        if line:
            lines.append(
                line
            )

    return "\n".join(
        lines
    ).strip()


def _parse_docx(
    *,
    path: Path,
    question: CompetitionQuestion,
    source: CompetitionSourceRecord,
    attachments_root: Path,
) -> CompetitionTextDocument:
    knowledge_source = (
        build_competition_knowledge_source(
            question=question,
            source=source,
            attachments_root=attachments_root,
        )
    )

    try:
        document = Document(
            path
        )

    except Exception as exc:
        raise CompetitionDocxOpenError(
            f"DOCX 无法打开: {path}"
        ) from exc

    blocks: list[
        CompetitionTextBlock
    ] = []

    # ========================================================
    # 注意：
    #
    # paragraph_index / table_index
    # 表示它们分别在原始 Word 文档中的顺序，
    # 而不是 Retrieval Block 顺序。
    #
    # block_index:
    #     所有实际输出 Block 的统一连续序号。
    # ========================================================

    paragraph_index = 0
    table_index = 0

    for item in (
        document.iter_inner_content()
    ):
        # ====================================================
        # Paragraph
        # ====================================================

        if isinstance(
            item,
            Paragraph,
        ):
            current_paragraph_index = (
                paragraph_index
            )

            paragraph_index += 1

            text = (
                _normalize_word_text(
                    item.text
                )
            )

            # 空段落本身没有检索价值，
            # 但 paragraph_index 仍然递增，
            # 从而保持原始位置稳定。
            if not text:
                continue

            block_index = len(
                blocks
            )

            style_name = (
                _paragraph_style_name(
                    item
                )
            )

            outline_level = (
                _effective_outline_level(
                    item
                )
            )

            blocks.append(
                CompetitionTextBlock(
                    block_id=_build_block_id(
                        doc_id=(
                            knowledge_source.doc_id
                        ),
                        block_index=(
                            block_index
                        ),
                    ),
                    source_id=(
                        knowledge_source.source_id
                    ),
                    doc_id=(
                        knowledge_source.doc_id
                    ),
                    source_type="word",
                    block_index=block_index,
                    block_type="paragraph",
                    text=text,
                    paragraph_index=(
                        current_paragraph_index
                    ),

                    # Word structure metadata
                    style_name=style_name,
                    outline_level=outline_level,
                )
            )

            continue

        # ====================================================
        # Table
        # ====================================================

        if isinstance(
            item,
            Table,
        ):
            current_table_index = (
                table_index
            )

            table_index += 1

            rows = _extract_table_rows(
                item
            )

            if not rows:
                continue

            text = _table_rows_to_text(
                rows
            )

            # 全空表格不产生检索 Block。
            if not text:
                continue

            block_index = len(
                blocks
            )

            blocks.append(
                CompetitionTextBlock(
                    block_id=_build_block_id(
                        doc_id=(
                            knowledge_source
                            .doc_id
                        ),
                        block_index=(
                            block_index
                        ),
                    ),
                    source_id=(
                        knowledge_source
                        .source_id
                    ),
                    doc_id=(
                        knowledge_source
                        .doc_id
                    ),
                    source_type="word",
                    block_index=(
                        block_index
                    ),
                    block_type="table",
                    text=text,
                    table_index=(
                        current_table_index
                    ),
                    table_rows=rows,
                )
            )

    if not blocks:
        raise (
            CompetitionEmptyTextDocumentError(
                "DOCX 没有产生任何"
                "可检索文本 Block"
            )
        )

    return CompetitionTextDocument(
        source=knowledge_source,
        blocks=tuple(
            blocks
        ),
    )

def _parse_legacy_doc(
    *,
    path: Path,
    question: CompetitionQuestion,
    source: CompetitionSourceRecord,
    attachments_root: Path,
) -> CompetitionTextDocument:
    """Convert one legacy DOC into a temporary DOCX and parse it."""

    with TemporaryDirectory(prefix="competition-legacy-doc-") as temp_directory:
        converted_path = convert_legacy_doc_to_docx(
            source_path=path,
            output_directory=Path(temp_directory),
        )

        return _parse_docx(
            path=converted_path,
            question=question,
            source=source,
            attachments_root=attachments_root,
        )

# ============================================================
# Public API
# ============================================================


def parse_competition_text_document(
    *,
    question: CompetitionQuestion,
    source: CompetitionSourceRecord,
    attachments_root: Path,
) -> CompetitionTextDocument:
    """
    Competition PDF / Word 统一 Parser 入口。

    当前支持：
        PDF
        DOCX

    Parse one competition PDF, DOCX, or legacy DOC document.
    """

    path = (
        resolve_competition_source_path(
            attachments_root=(
                attachments_root
            ),
            source=source,
        )
    )

    extension = (
        source.extension.casefold()
    )

    if (
        source.source_type == "pdf"
        and extension == ".pdf"
    ):
        return _parse_pdf(
            path=path,
            question=question,
            source=source,
            attachments_root=(
                attachments_root
            ),
        )

    if (
        source.source_type == "word"
        and extension == ".docx"
    ):
        return _parse_docx(
            path=path,
            question=question,
            source=source,
            attachments_root=(
                attachments_root
            ),
        )

    if path.suffix.lower() == ".doc":
        return _parse_legacy_doc(
            path=path,
            question=question,
            source=source,
            attachments_root=attachments_root,
        )
    raise (
        CompetitionUnsupportedTextFormatError(
            "不支持的 Competition "
            "文本格式: "
            f"source_type="
            f"{source.source_type}, "
            f"extension={extension}"
        )
    )